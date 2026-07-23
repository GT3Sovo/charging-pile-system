# -*- coding: utf-8 -*-
"""最终验证图片嵌入结果"""

import sys, io, re
from docx import Document
from lxml import etree

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

doc = Document(r"C:\Users\15479\Desktop\软工 真最终版\ChargingPileDispatchBillingSystem\第2次作业 G12组 (1).docx")

# 统计图片总数
body_xml = etree.tostring(doc.element.body, encoding='unicode')
inlines = re.findall(r'<wp:inline[\s\S]*?</wp:inline>', body_xml)
print(f"文档总图片数: {len(inlines)} (原22 + 新增9 = 预期31)")

# 验证关键位置
checks = ['图 5 指令reportBreakdown', '图 8 指令getDailyReport', 
          '图 11 指令resetSimulation', '图 13 指令nextEvent']
for i, p in enumerate(doc.paragraphs):
    for chk in checks:
        if chk in p.text:
            nxml = etree.tostring(doc.paragraphs[i+1]._element, encoding='unicode')
            has_img = 'wp:inline' in nxml
            status = "有图片" if has_img else "无图片!!!"
            print(f"段落 {i} [{chk[:30]}...] 后: {status}")

# 检查所有新caption后面是否有图片
print("\n=== 检查所有9个新图 ===")
new_captions = [
    '图 5 指令reportBreakdown',
    '图 6 指令triggerRecovery',
    '图 7 指令generateBill',
    '图 8 指令getDailyReport',
    '图 9 指令getReportSummary',
    '图 10 指令exportCsv',
    '图 11 指令resetSimulation',
    '图 12 指令setSimulationMode',
    '图 13 指令nextEvent',
]
all_ok = True
for i, p in enumerate(doc.paragraphs):
    for chk in new_captions:
        if chk in p.text:
            nxml = etree.tostring(doc.paragraphs[i+1]._element, encoding='unicode')
            if 'wp:inline' not in nxml:
                print(f"  MISSING: 段落 {i} [{chk[:40]}]")
                all_ok = False

if all_ok:
    print("  全部9个图后面都有图片！")